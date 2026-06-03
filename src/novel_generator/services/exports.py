from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ..models import Artifact, ChapterDraft, GenerationRun, Project
from .prompts import sanitize_chapter_content


@dataclass(frozen=True)
class ExportProfile:
    id: str
    label: str
    output_format: str
    filename: str
    kind: str
    description: str
    page_width_inches: float | None = None
    page_height_inches: float | None = None
    top_margin_inches: float = 0.75
    bottom_margin_inches: float = 0.75
    inner_margin_inches: float = 0.75
    outer_margin_inches: float = 0.75
    include_front_matter: bool = False
    chapter_page_breaks: bool = False
    publication_ready: bool = False


@dataclass(frozen=True)
class PublicationFrontMatter:
    author_name: str
    copyright_year: str
    publisher: str
    dedication: str
    author_note: str
    isbn: str = ""
    ai_disclosure: str = ""


EXPORT_PROFILES = {
    "draft_markdown": ExportProfile(
        id="draft_markdown",
        label="Draft Markdown",
        output_format="markdown",
        filename="manuscript.md",
        kind="markdown",
        description="Review-oriented Markdown draft.",
    ),
    "draft_docx": ExportProfile(
        id="draft_docx",
        label="Draft DOCX",
        output_format="docx",
        filename="manuscript.docx",
        kind="docx",
        description="Review-oriented DOCX draft.",
    ),
    "ebook_markdown": ExportProfile(
        id="ebook_markdown",
        label="Ebook Markdown",
        output_format="markdown",
        filename="publication-ebook.md",
        kind="publication-markdown",
        description="Markdown manuscript layout helper with supplied front matter.",
        include_front_matter=True,
        publication_ready=True,
    ),
    "ebook_docx": ExportProfile(
        id="ebook_docx",
        label="Ebook DOCX",
        output_format="docx",
        filename="publication-ebook.docx",
        kind="publication-docx",
        description="DOCX manuscript layout helper with supplied front matter and chapter page starts.",
        include_front_matter=True,
        chapter_page_breaks=True,
        publication_ready=True,
    ),
    "print_5x8": ExportProfile(
        id="print_5x8",
        label='Print 5" x 8"',
        output_format="docx",
        filename="publication-print-5x8.docx",
        kind="publication-docx",
        description='Print layout helper interior sized to 5" x 8".',
        page_width_inches=5,
        page_height_inches=8,
        inner_margin_inches=0.75,
        outer_margin_inches=0.6,
        include_front_matter=True,
        chapter_page_breaks=True,
        publication_ready=True,
    ),
    "print_5_5x8_5": ExportProfile(
        id="print_5_5x8_5",
        label='Print 5.5" x 8.5"',
        output_format="docx",
        filename="publication-print-5-5x8-5.docx",
        kind="publication-docx",
        description='Print layout helper interior sized to 5.5" x 8.5".',
        page_width_inches=5.5,
        page_height_inches=8.5,
        include_front_matter=True,
        chapter_page_breaks=True,
        publication_ready=True,
    ),
    "print_6x9": ExportProfile(
        id="print_6x9",
        label='Print 6" x 9"',
        output_format="docx",
        filename="publication-print-6x9.docx",
        kind="publication-docx",
        description='Print layout helper interior sized to 6" x 9".',
        page_width_inches=6,
        page_height_inches=9,
        top_margin_inches=0.8,
        bottom_margin_inches=0.8,
        include_front_matter=True,
        chapter_page_breaks=True,
        publication_ready=True,
    ),
    "print_a5": ExportProfile(
        id="print_a5",
        label='Print A5 / 5.83" x 8.27"',
        output_format="docx",
        filename="publication-print-a5.docx",
        kind="publication-docx",
        description='Print layout helper interior sized to A5 / 5.83" x 8.27".',
        page_width_inches=5.83,
        page_height_inches=8.27,
        include_front_matter=True,
        chapter_page_breaks=True,
        publication_ready=True,
    ),
}


PUBLICATION_PROFILE_IDS = ["ebook_markdown", "ebook_docx", "print_5x8", "print_5_5x8_5", "print_6x9", "print_a5"]


def publication_export_options() -> list[ExportProfile]:
    return [EXPORT_PROFILES[profile_id] for profile_id in PUBLICATION_PROFILE_IDS]


def export_profile(profile_id: str) -> ExportProfile:
    try:
        return EXPORT_PROFILES[profile_id]
    except KeyError as exc:
        raise ValueError("Choose a supported export profile.") from exc


def _normalize_front_matter(front_matter: PublicationFrontMatter | dict[str, str] | None) -> PublicationFrontMatter:
    if front_matter is None:
        raise ValueError("Publication front matter is required for publication exports.")
    if isinstance(front_matter, PublicationFrontMatter):
        payload = front_matter
    elif isinstance(front_matter, dict):
        payload = PublicationFrontMatter(
            author_name=str(front_matter.get("author_name", "") or "").strip(),
            copyright_year=str(front_matter.get("copyright_year", "") or "").strip(),
            publisher=str(front_matter.get("publisher", "") or "").strip(),
            dedication=str(front_matter.get("dedication", "") or "").strip(),
            author_note=str(front_matter.get("author_note", "") or "").strip(),
            isbn=str(front_matter.get("isbn", "") or "").strip(),
            ai_disclosure=str(front_matter.get("ai_disclosure", "") or "").strip(),
        )
    else:
        raise ValueError("Publication front matter must be a supported metadata object.")

    required = {
        "author name": payload.author_name,
        "copyright year": payload.copyright_year,
        "publisher or imprint": payload.publisher,
        "dedication": payload.dedication,
        "author note": payload.author_note,
    }
    missing = [label for label, value in required.items() if not value]
    if missing:
        raise ValueError("Publication front matter is missing: " + ", ".join(missing) + ".")

    placeholder_values = {
        label: value
        for label, value in {
            **required,
            "isbn": payload.isbn,
            "AI disclosure": payload.ai_disclosure,
        }.items()
        if re.search(r"\[[^\]]+\]", value)
    }
    if placeholder_values:
        raise ValueError(
            "Publication front matter cannot contain bracketed placeholders: "
            + ", ".join(placeholder_values.keys())
            + "."
        )
    return payload


def export_run_artifacts(
    artifacts_dir: Path,
    project: Project,
    run: GenerationRun,
    chapters: list[ChapterDraft],
    qa_report_markdown: str | None = None,
    developmental_rewrite_markdown: str | None = None,
    revised_outline_markdown: str | None = None,
    developmental_qa_markdown: str | None = None,
) -> list[Artifact]:
    run_dir = artifacts_dir / run.id
    run_dir.mkdir(parents=True, exist_ok=True)

    markdown_name = "manuscript.md"
    markdown_path = run_dir / markdown_name
    markdown_path.write_text(render_markdown(project, chapters), encoding="utf-8")

    docx_name = "manuscript.docx"
    docx_path = run_dir / docx_name
    render_docx(project, chapters, docx_path)

    artifacts = [
        Artifact(
            kind="markdown",
            filename=markdown_name,
            relative_path=str(markdown_path.relative_to(artifacts_dir)),
            content_type="text/markdown",
        ),
        Artifact(
            kind="docx",
            filename=docx_name,
            relative_path=str(docx_path.relative_to(artifacts_dir)),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    ]
    if qa_report_markdown is not None:
        qa_name = "qa-report.md"
        qa_path = run_dir / qa_name
        qa_path.write_text(qa_report_markdown, encoding="utf-8")
        artifacts.append(
            Artifact(
                kind="qa-report",
                filename=qa_name,
                relative_path=str(qa_path.relative_to(artifacts_dir)),
                content_type="text/markdown",
            )
        )
    if developmental_rewrite_markdown is not None:
        rewrite_name = "developmental-rewrite-report.md"
        rewrite_path = run_dir / rewrite_name
        rewrite_path.write_text(developmental_rewrite_markdown, encoding="utf-8")
        artifacts.append(
            Artifact(
                kind="developmental-rewrite-report",
                filename=rewrite_name,
                relative_path=str(rewrite_path.relative_to(artifacts_dir)),
                content_type="text/markdown",
            )
        )
    if revised_outline_markdown is not None:
        outline_name = "revised-outline.md"
        outline_path = run_dir / outline_name
        outline_path.write_text(revised_outline_markdown, encoding="utf-8")
        artifacts.append(
            Artifact(
                kind="revised-outline",
                filename=outline_name,
                relative_path=str(outline_path.relative_to(artifacts_dir)),
                content_type="text/markdown",
            )
        )
    if developmental_qa_markdown is not None:
        qa_comparison_name = "developmental-qa-comparison.md"
        qa_comparison_path = run_dir / qa_comparison_name
        qa_comparison_path.write_text(developmental_qa_markdown, encoding="utf-8")
        artifacts.append(
            Artifact(
                kind="developmental-qa-report",
                filename=qa_comparison_name,
                relative_path=str(qa_comparison_path.relative_to(artifacts_dir)),
                content_type="text/markdown",
            )
        )
    return artifacts


def export_publication_artifact(
    artifacts_dir: Path,
    project: Project,
    run: GenerationRun,
    chapters: list[ChapterDraft],
    profile_id: str,
    *,
    include_ai_disclosure: bool = False,
    front_matter: PublicationFrontMatter | dict[str, str] | None = None,
) -> Artifact:
    profile = export_profile(profile_id)
    if not profile.publication_ready:
        raise ValueError("Choose a publication export profile.")
    publication_front_matter = _normalize_front_matter(front_matter) if profile.include_front_matter else None

    run_dir = artifacts_dir / run.id
    run_dir.mkdir(parents=True, exist_ok=True)
    destination = run_dir / profile.filename

    if profile.output_format == "markdown":
        destination.write_text(
            render_publication_markdown(
                project,
                chapters,
                profile,
                include_ai_disclosure=include_ai_disclosure,
                front_matter=publication_front_matter,
            ),
            encoding="utf-8",
        )
        content_type = "text/markdown"
    elif profile.output_format == "docx":
        render_publication_docx(
            project,
            chapters,
            destination,
            profile,
            include_ai_disclosure=include_ai_disclosure,
            front_matter=publication_front_matter,
        )
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise ValueError("Unsupported export profile format.")

    return Artifact(
        kind=profile.kind,
        filename=profile.filename,
        relative_path=str(destination.relative_to(artifacts_dir)),
        content_type=content_type,
    )


def render_markdown(project: Project, chapters: list[ChapterDraft]) -> str:
    lines = [f"# {project.title}", "", project.premise, ""]
    for chapter in chapters:
        content = sanitize_chapter_content(chapter.content or "")
        lines.extend(
            [
                f"## Chapter {chapter.chapter_number}: {chapter.title}",
                "",
                content,
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def render_docx(project: Project, chapters: list[ChapterDraft], destination: Path) -> None:
    document = Document()
    document.add_heading(project.title, level=0)
    document.add_paragraph(project.premise)
    for chapter in chapters:
        content = sanitize_chapter_content(chapter.content or "")
        document.add_heading(f"Chapter {chapter.chapter_number}: {chapter.title}", level=1)
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
    document.save(destination)


def render_publication_markdown(
    project: Project,
    chapters: list[ChapterDraft],
    profile: ExportProfile,
    *,
    include_ai_disclosure: bool = False,
    front_matter: PublicationFrontMatter | None = None,
) -> str:
    lines: list[str] = [f"# {project.title}", ""]
    if profile.include_front_matter:
        if front_matter is None:
            raise ValueError("Publication front matter is required for publication exports.")
        lines.extend(_front_matter_markdown(project, front_matter, include_ai_disclosure=include_ai_disclosure))
    for chapter in chapters:
        lines.extend(
            [
                f"## Chapter {chapter.chapter_number}: {chapter.title}",
                "",
                _clean_publication_text(chapter.content or ""),
                "",
            ]
        )
    return "\n".join(line.rstrip() for line in lines).strip() + "\n"


def render_publication_docx(
    project: Project,
    chapters: list[ChapterDraft],
    destination: Path,
    profile: ExportProfile,
    *,
    include_ai_disclosure: bool = False,
    front_matter: PublicationFrontMatter | None = None,
) -> None:
    document = Document()
    _configure_publication_document(document, profile)
    _add_title_page(document, project, front_matter)
    if profile.include_front_matter:
        if front_matter is None:
            raise ValueError("Publication front matter is required for publication exports.")
        _add_front_matter(document, project, front_matter, include_ai_disclosure=include_ai_disclosure)
    for chapter in chapters:
        if profile.chapter_page_breaks:
            document.add_page_break()
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.style = document.styles["Heading 1"]
        heading.add_run(f"Chapter {chapter.chapter_number}: {chapter.title}")
        for paragraph in _publication_paragraphs(chapter.content or ""):
            document.add_paragraph(paragraph)
    document.save(destination)


def _configure_publication_document(document: Document, profile: ExportProfile) -> None:
    section = document.sections[0]
    if profile.page_width_inches and profile.page_height_inches:
        section.page_width = Inches(profile.page_width_inches)
        section.page_height = Inches(profile.page_height_inches)
    section.top_margin = Inches(profile.top_margin_inches)
    section.bottom_margin = Inches(profile.bottom_margin_inches)
    section.left_margin = Inches(profile.inner_margin_inches)
    section.right_margin = Inches(profile.outer_margin_inches)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.25)

    heading = document.styles["Heading 1"]
    heading.font.name = "Times New Roman"
    heading.font.size = Pt(16)
    heading.paragraph_format.space_before = Pt(24)
    heading.paragraph_format.space_after = Pt(18)
    heading.paragraph_format.first_line_indent = None


def _add_title_page(document: Document, project: Project, front_matter: PublicationFrontMatter | None = None) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(180)
    title.paragraph_format.first_line_indent = None
    title_run = title.add_run(project.title)
    title_run.bold = True
    title_run.font.size = Pt(20)

    byline = document.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.first_line_indent = None
    byline.add_run(f"by {front_matter.author_name if front_matter else 'Author'}")


def _add_front_matter(
    document: Document,
    project: Project,
    front_matter: PublicationFrontMatter,
    *,
    include_ai_disclosure: bool,
) -> None:
    copyright_lines = [
        f"{project.title}",
        f"Copyright (c) {front_matter.copyright_year} {front_matter.author_name}. All rights reserved.",
        f"Publisher: {front_matter.publisher}",
    ]
    if front_matter.isbn:
        copyright_lines.append(f"ISBN: {front_matter.isbn}")
    _add_front_matter_page(document, "Copyright", copyright_lines)
    _add_front_matter_page(document, "Dedication", [front_matter.dedication])
    _add_front_matter_page(document, "Author Note", [front_matter.author_note])
    if include_ai_disclosure:
        disclosure = (
            front_matter.ai_disclosure
            or "This manuscript was drafted and edited with AI assistance under human direction."
        )
        _add_front_matter_page(
            document,
            "AI-Assisted Disclosure",
            [disclosure],
        )


def _add_front_matter_page(document: Document, heading_text: str, paragraphs: list[str]) -> None:
    document.add_page_break()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = None
    run = heading.add_run(heading_text)
    run.bold = True
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.first_line_indent = None


def _front_matter_markdown(
    project: Project,
    front_matter: PublicationFrontMatter,
    *,
    include_ai_disclosure: bool,
) -> list[str]:
    lines = [
        "## Copyright",
        "",
        f"{project.title}",
        "",
        f"Copyright (c) {front_matter.copyright_year} {front_matter.author_name}. All rights reserved.",
        "",
        f"Publisher: {front_matter.publisher}",
        "",
        "## Dedication",
        "",
        front_matter.dedication,
        "",
        "## Author Note",
        "",
        front_matter.author_note,
        "",
    ]
    if front_matter.isbn:
        lines[8:8] = [f"ISBN: {front_matter.isbn}", ""]
    if include_ai_disclosure:
        lines.extend(
            [
                "## AI-Assisted Disclosure",
                "",
                front_matter.ai_disclosure
                or "This manuscript was drafted and edited with AI assistance under human direction.",
                "",
            ]
        )
    return lines


def _publication_paragraphs(content: str) -> list[str]:
    text = _clean_publication_text(content)
    return [paragraph.strip() for paragraph in re.split(r"\n\s*\n+", text) if paragraph.strip()]


def _clean_publication_text(content: str) -> str:
    cleaned = sanitize_chapter_content(content)
    lines: list[str] = []
    for line in cleaned.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            continue
        if re.match(r"^#{1,6}\s*chapter\s+\d+\b", stripped, flags=re.IGNORECASE):
            continue
        stripped = re.sub(r"^#{1,6}\s+", "", stripped)
        if re.match(r"^chapter\s+\d+\b", stripped, flags=re.IGNORECASE) and len(stripped) <= 80:
            continue
        stripped = re.sub(r"^\s*[-*+]\s+", "", stripped)
        stripped = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        stripped = re.sub(r"__(.*?)__", r"\1", stripped)
        stripped = stripped.replace("`", "")
        lines.append(stripped)
    return "\n".join(lines).strip()
